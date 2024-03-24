from docx import Document
from docx.shared import Inches

document = Document()

#  profile picture
document.add_picture(
    'profile-pic.png',
     width=Inches(2.0)
)

# user details; name, phone and email
name = input('What is your name? ')
last_name = input('Put in your last name')
phone_number = input('what is your phone number')
email = input('We also need your email address')

document.add_paragraph(
    name + last_name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading( 'About me' )
about_me = input('Tell us about yourself? ')
document.add_paragraph(about_me)

# Work Experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company')
from_date= input('From Date')
to_date= input('To Date')


p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True
experience_details = input(
    'Describe your experience to us at' + company)
p.add_run(experience_details)


document.save('cv.docx')